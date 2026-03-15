import os
import pandas as pd
from docx import Document
from docx.shared import Pt

def migrate_logbook():
    # Paths
    excel_path = "/Users/brianrizqi/Documents/Post Doc/agentic-ai-native-ads/Logbook_PostDoc_AI_Native_Ads.xlsx"
    word_path = "/Users/brianrizqi/Documents/Post Doc/agentic-ai-native-ads/Catatan-Harian-Top-Tiers-Profesor-Doktor.docx"
    output_path = "/Users/brianrizqi/Documents/Post Doc/agentic-ai-native-ads/Catatan-Harian-Top-Tiers-Profesor-Doktor_Updated.docx"

    # Read Excel data
    print(f"Reading data from {excel_path}...")
    df = pd.read_excel(excel_path)
    
    # Expected columns: No, Tanggal, Ringkasan Aktivitas Penelitian
    # Rename columns if necessary for consistency
    df.columns = ['No', 'Tanggal', 'Kegiatan']

    # Load Word document
    print(f"Opening Word document {word_path}...")
    doc = Document(word_path)
    
    if not doc.tables:
        print("No tables found in the Word document.")
        return

    table = doc.tables[0]
    
    # Remove existing placeholder rows (starting from row 1)
    # The table has 5 rows: Header (0), Placeholder 1, 2, 3, Footer (4)
    # We want to keep the header and replace everything else with real data
    
    # It's safer to delete from the end to avoid index shifts
    for _ in range(len(table.rows) - 1):
        row = table.rows[1]
        parent = row._element.getparent()
        parent.remove(row._element)

    # Add new rows from Excel data
    print(f"Adding {len(df)} rows to the table...")
    for index, row_data in df.iterrows():
        new_row = table.add_row().cells
        new_row[0].text = str(row_data['No'])
        new_row[1].text = str(row_data['Tanggal'])
        
        # Format the "Kegiatan" cell
        activity = str(row_data['Kegiatan'])
        # Use existing formatting style: Catatan: ...\nDokumen Pendukung:
        new_row[2].text = f"Catatan: {activity}\nDokumen Pendukung:"

    # Save the updated document
    print(f"Saving updated document to {output_path}...")
    doc.save(output_path)
    print("Migration completed successfully!")

if __name__ == "__main__":
    migrate_logbook()
